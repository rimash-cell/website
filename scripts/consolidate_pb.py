import os
import sys
from docx import Document

def get_text_from_docx(file_path):
    doc = Document(file_path)
    return [para.text for para in doc.paragraphs]

def add_section(master_doc, title, content_paragraphs):
    master_doc.add_heading(title, level=1)
    for para in content_paragraphs:
        if para.strip():
            master_doc.add_paragraph(para)
        else:
            master_doc.add_paragraph("")

def main():
    raw_dir = r"c:\Users\rimaa\3D Objects\Website\Rima's - personal_brand_raw"
    output_path = os.path.join(raw_dir, "AfterSense_Master_OS_v1.docx")
    
    master_doc = Document()
    master_doc.add_heading("AfterSense Master OS", 0)
    master_doc.add_paragraph("This document serves as the unified, canonical source of truth for the AfterSense Brand and Personal Brand System.")
    
    # 1. Foundation & Master Document Consolidation
    master_path = os.path.join(raw_dir, "Master Brand Playbook (Canonical v1).docx")
    if os.path.exists(master_path):
        print(f"Adding {master_path}...")
        add_section(master_doc, "I. Brand Foundation & Playbook", get_text_from_docx(master_path))

    # 2. Voice & Messaging
    voice_path = os.path.join(raw_dir, "Aftersense - Brand Voice & Messaging Rules.docx")
    if os.path.exists(voice_path):
        print(f"Adding {voice_path}...")
        add_section(master_doc, "II. Voice & Messaging Rules", get_text_from_docx(voice_path))

    # 3. Arabic Voice Addendum
    arabic_path = os.path.join(raw_dir, "Arabic_Brand_Voice_and_Language_Rules_Addendum_v1.1.docx")
    if os.path.exists(arabic_path):
        print(f"Adding {arabic_path}...")
        add_section(master_doc, "III. Arabic Language & Cultural Codes", get_text_from_docx(arabic_path))

    # 4. Constraints & Guardrails
    guardrails_path = os.path.join(raw_dir, "Constraints & Guardrails.docx")
    if os.path.exists(guardrails_path):
        print(f"Adding {guardrails_path}...")
        add_section(master_doc, "IV. Constraints & Guardrails", get_text_from_docx(guardrails_path))

    # 5. Service & Offer Architecture
    service_path = os.path.join(raw_dir, "Service & Offer Architecture.docx")
    if os.path.exists(service_path):
        print(f"Adding {service_path}...")
        add_section(master_doc, "V. Service & Offer Architecture", get_text_from_docx(service_path))

    # 6. SOPs
    sop_path = os.path.join(raw_dir, "Execution SOPs & Playbooks.docx")
    if os.path.exists(sop_path):
        print(f"Adding {sop_path}...")
        add_section(master_doc, "VI. Execution SOPs & Playbooks", get_text_from_docx(sop_path))

    master_doc.save(output_path)
    print(f"Successfully created: {output_path}")

if __name__ == "__main__":
    main()
